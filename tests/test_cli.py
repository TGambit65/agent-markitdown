from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from agent_markitdown.converter import default_output_path, validate_local_input  # noqa: E402


def _run(args: list[str]) -> subprocess.CompletedProcess[str]:
    env = os.environ.copy()
    env["PYTHONPATH"] = str(SRC)
    return subprocess.run(
        [sys.executable, "-m", "agent_markitdown.cli", *args],
        cwd=ROOT,
        env=env,
        text=True,
        capture_output=True,
        check=False,
    )


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Test Doc", level=1)
    doc.add_paragraph("Hello from DOCX.")
    doc.save(path)


def _make_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path))
    c.drawString(72, 720, "Hello from PDF.")
    c.save()


def test_validate_local_input_rejects_unsupported(tmp_path: Path) -> None:
    sample = tmp_path / "archive.zip"
    sample.write_bytes(b"PK\x03\x04")
    try:
        validate_local_input(sample)
    except ValueError as exc:
        assert "Unsupported extension" in str(exc)
    else:
        raise AssertionError("expected ValueError")


def test_default_output_path(tmp_path: Path) -> None:
    sample = tmp_path / "report.pdf"
    assert default_output_path(sample) == tmp_path / "report.pdf.md"


def test_convert_docx_to_json(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    _make_docx(source)
    result = _run(["convert", str(source), "--json"])
    assert result.returncode == 0, result.stderr
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["extension"] == ".docx"
    assert "Hello from DOCX." in payload["markdown"]


def test_convert_pdf_sidecar(tmp_path: Path) -> None:
    source = tmp_path / "sample.pdf"
    _make_pdf(source)
    result = _run(["convert", str(source), "--sidecar"])
    assert result.returncode == 0, result.stderr
    output_path = Path(result.stdout.strip())
    assert output_path.exists()
    text = output_path.read_text(encoding="utf-8")
    assert "Hello from PDF." in text


def test_review_pack(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    _make_docx(source)
    result = _run(["review-pack", str(source)])
    assert result.returncode == 0, result.stderr
    assert "# Agent Review Pack" in result.stdout
    assert "Hello from DOCX." in result.stdout


def test_doctor() -> None:
    result = _run(["doctor"])
    assert result.returncode == 0
    payload = json.loads(result.stdout)
    assert payload["local_only"] is True
    assert "review-pack" in payload["subcommands"]
